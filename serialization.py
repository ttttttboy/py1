# coding:utf-8
#  ref: https://zhuanlan.zhihu.com/p/21716087
import time
from docx import Document
from config import *



def demo_doc():
    item = [u"这是标题", u"国发办[2017]45号", u"这是正文第一行\n这是正文第二行\n这是结尾"]
    demo_path = r"C:\Users\Leon\Documents\Files\GitHub\py1\demo.docx"

    ddoc = Document()
    ddoc.add_paragraph(item[0], 'Title')
    ddoc.add_paragraph(item[1], 'Subtitle')
    ddoc.add_paragraph(item[2], 'Normal')

    ddoc.save(demo_path)


def out2File_docx(buffer, item):
    path_full_file = PATH_OUTPUT_FOLDER + os.path.sep + item[1] + item[0][0:128] + ".txt"
    with open(path_full_file, 'w', encoding='utf-8') as resf:
        resf.write(buffer)
        print(" OK! to file %s" % item[0])


# normal 35.5
def out2File_txt(buffer, item):
    path_full_file = PATH_OUTPUT_FOLDER + os.path.sep + item[1] + item[0][0:128] + ".txt"
    with open(path_full_file, 'w',
              encoding='utf-8') as resf:
        # resf.write(buffer)
        time.sleep(8)
        print("        #%s serialised. %s" % (item[3], item[0]))


